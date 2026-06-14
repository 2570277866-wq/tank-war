#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成服务端报告书 docx"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = '宋体'
style.paragraph_format.line_spacing = 1.5

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run('课程设计报告书')
run.font.size = Pt(26)
run.font.bold = True
run.font.name = '黑体'

doc.add_paragraph('')

info_items = [
    '题目：坦克大战',
    '学  院     计算机科学与工程学院',
    '专  业            25计类',
    '学生姓名           ___________',
    '学生学号          ___________',
    '指导教师            沃焱',
    '课程编号          045101571',
    '课程学分             2',
    '起始日期  2026年3月至2026年6月',
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(14)
    run.font.name = '宋体'

doc.add_page_break()

# ============================================================
# TITLE
# ============================================================
h = doc.add_heading('坦克大战', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h.runs:
    run.font.name = '黑体'
doc.add_paragraph('')

# ============================================================
# 一、选题背景
# ============================================================
doc.add_heading('一、选题背景', level=2)

doc.add_heading('1.1 课题来源', level=3)
doc.add_paragraph(
    '坦克大战（Battle City）是经典的FC平台游戏，其核心玩法——坦克对战、障碍物策略、'
    '实时对抗——至今仍具有很高的娱乐价值和设计参考意义。随着网络技术的发展，将经典单机'
    '玩法扩展为联机对战模式，不仅能够还原经典体验，还能增加玩家间的实时互动与竞技性。'
)
doc.add_paragraph(
    '本课题选择"双人联机坦克大战"作为课程设计题目，旨在综合运用C++面向对象编程、'
    'Windows网络编程、多线程同步、图形界面开发等技术，完成一个具备完整游戏流程、'
    '网络联机、服务端权威判定的实时对战游戏。'
)

doc.add_heading('1.2 选题意义', level=3)
doc.add_paragraph(
    '1. 技术综合性强：项目涵盖C++类体系设计、WinSock2 TCP网络通信、多线程并发控制、'
    'EasyX图形渲染等多个技术方向，是对课程所学知识的综合实践。'
)
doc.add_paragraph(
    '2. 架构设计典型：采用客户端-服务端（C/S）架构，服务端权威判定、客户端预测与插值，'
    '是网络游戏开发的经典模式，具有工程参考价值。'
)
doc.add_paragraph(
    '3. 实时同步挑战：20Hz状态快照广播、客户端插值平滑、断线重连等机制的实现，'
    '涉及实时系统设计的核心问题。'
)
doc.add_paragraph(
    '4. 团队协作实践：双人分工开发、协议联调、Git版本管理，模拟真实项目协作流程。'
)

doc.add_heading('1.3 国内外研究现状', level=3)
doc.add_paragraph(
    '实时对战游戏在网络游戏领域已有成熟方案，如帧同步（Lockstep）和状态同步'
    '（State Synchronization）两种主流架构。本项目采用状态同步方案，服务端以20Hz'
    '频率广播游戏快照，客户端进行插值渲染，在局域网环境下可达到流畅的对战体验。'
    '该方案实现复杂度适中，适合课程设计规模。'
)

# ============================================================
# 二、系统设计
# ============================================================
doc.add_heading('二、系统设计', level=2)

doc.add_heading('2.1 项目目录结构', level=3)
dir_text = (
    'TankWar/（本人负责部分）\n'
    '├── Server/\n'
    '│   ├── main.cpp              # 服务端入口\n'
    '│   ├── Net/                  # 网络层\n'
    '│   │   ├── TCPServer.cpp/h   # TCP监听、Accept、Tick调度\n'
    '│   │   └── Session.cpp/h     # 客户端会话管理\n'
    '│   ├── Game/                 # 游戏逻辑层\n'
    '│   │   ├── Room.cpp/h        # 房间管理、匹配、断线重连\n'
    '│   │   ├── GameWorld.cpp/h   # 游戏世界Tick、权威判定\n'
    '│   │   └── Judge.h           # 碰撞检测工具函数\n'
    '│   ├── Data/                 # 数据持久化层\n'
    '│   │   ├── UserManager.cpp/h # 用户注册登录\n'
    '│   │   ├── RecordManager.cpp/h # 对局记录\n'
    '│   │   └── Leaderboard.cpp/h # 排行榜\n'
    '│   └── Core/                 # 基础设施\n'
    '│       ├── Clock.h           # 高精度计时器\n'
    '│       └── Logger.h          # 日志系统\n'
    '└── Common/                   # 共享协议（双方共同维护）\n'
    '    ├── Protocol.h            # 消息ID、结构体定义\n'
    '    └── Config.h              # 全局常量、地图配置'
)
p = doc.add_paragraph()
run = p.add_run(dir_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('2.2 系统总体架构', level=3)
doc.add_paragraph(
    '系统采用C/S架构，分为客户端和服务端两大模块：'
)
doc.add_paragraph(
    '客户端A（EasyX渲染） ←——TCP 9527端口——→ 服务端（权威判定） ←——TCP 9527端口——→ 客户端B（EasyX渲染）'
)
doc.add_paragraph('核心设计原则：')
doc.add_paragraph(
    '1. 服务端权威：所有伤害判定、碰撞检测由服务端执行，客户端仅负责渲染和上报输入'
)
doc.add_paragraph(
    '2. 状态同步：服务端每50ms广播一次完整游戏快照，客户端按快照更新本地状态'
)
doc.add_paragraph(
    '3. 客户端预测：客户端在等待快照期间先本地预测移动，收到快照后修正'
)

doc.add_heading('2.3 功能模块划分', level=3)
doc.add_paragraph('本人负责服务端全部模块的开发，按功能分为以下四层：')

p = doc.add_paragraph()
p.add_run('服务端模块：').bold = True
doc.add_paragraph(
    '网络层(Net)：TCPServer（TCP监听、Accept循环、Tick线程）、Session（客户端会话、消息分发）'
)
doc.add_paragraph(
    '游戏层(Game)：Room（房间管理、匹配、断线重连）、GameWorld（游戏Tick、碰撞判定、快照打包）、'
    'Judge（碰撞工具函数）'
)
doc.add_paragraph(
    '数据层(Data)：UserManager（用户注册登录密码哈希）、RecordManager（对局记录保存）、'
    'Leaderboard（排行榜排序）'
)
doc.add_paragraph(
    '核心层(Core)：Clock（高精度计时器）、Logger（日志系统）'
)

doc.add_heading('2.4 坦克类体系设计', level=3)
doc.add_paragraph('采用工厂模式+多态设计：')
tank_text = (
    'Tank（抽象基类）\n'
    '├── HeavyTank — 重型坦克：HP:200 速度:慢 伤害:大 技能:铁壁护盾(3秒免伤, CD:15s)\n'
    '├── LightTank — 轻型坦克：HP:150 速度:中 伤害:中 技能:涡轮冲刺(2秒双倍速, CD:12s)\n'
    '└── ScoutTank — 侦察坦克：HP:100 速度:快 伤害:小 技能:弹幕散射(3颗扇形弹, CD:10s)'
)
p = doc.add_paragraph()
run = p.add_run(tank_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph(
    '基类定义纯虚接口（getDamage/getSpeed/getSkillType等），子类通过override返回不同属性值，'
    '工厂方法Tank::create()根据TankType创建对应子类实例。'
)

doc.add_heading('2.5 网络协议设计', level=3)
doc.add_paragraph('消息格式：[2字节 MsgID] [2字节 bodyLen] [body变长]')
doc.add_paragraph('消息分类：')
doc.add_paragraph('0x01xx — 认证（登录/注册），客户端↔服务端')
doc.add_paragraph('0x02xx — 房间（加入/离开/匹配/选坦克/重连），双向')
doc.add_paragraph('0x03xx — 游戏同步（输入/快照/射击/技能/命中），双向')
doc.add_paragraph('0x04xx — 结算（游戏结束/排行榜），双向')
doc.add_paragraph('0x05xx — 心跳，双向')
doc.add_paragraph('同步流程：')
doc.add_paragraph('客户端 → C2S_INPUT(按键状态) → 服务端')
doc.add_paragraph('服务端 → 逻辑Tick(50ms) → 广播S2C_SNAPSHOT → 客户端')
doc.add_paragraph('客户端收到快照 → 插值渲染')

doc.add_heading('2.6 地图系统设计', level=3)
doc.add_paragraph('砖墙：可破坏，HP=50，阻挡坦克和子弹，子弹命中减HP，HP=0消失')
doc.add_paragraph('钢墙：不可破坏，阻挡坦克和子弹，子弹命中消失墙不损')
doc.add_paragraph('草丛：隐身效果，不阻挡坦克，子弹穿过')
doc.add_paragraph('水域：不可通行，阻挡坦克，子弹穿过')

doc.add_heading('2.7 断线重连机制', level=3)
doc.add_paragraph('1. 服务端心跳检测：3秒无心跳标记为断线')
doc.add_paragraph('2. 断线后游戏暂停（RoomState::PAUSED），通知双方')
doc.add_paragraph('3. 10秒重连窗口期内，客户端发送C2S_RECONNECT可恢复')
doc.add_paragraph('4. 超时未重连，断线方判负')

# ============================================================
# 三、系统实现（服务端——本人负责部分）
# ============================================================
doc.add_heading('三、系统实现', level=2)

doc.add_paragraph(
    '本人在项目中负责服务端开发，共计实现7个功能模块。服务端是整个游戏的核心枢纽，'
    '承担客户端连接管理、用户认证、房间匹配、游戏逻辑权威判定、快照广播、战绩记录与排行榜等功能。'
)

# ---------- (一) Common 公共模块 ----------
doc.add_heading('（一）Common 公共模块', level=3)
doc.add_paragraph(
    '本模块为全局公共资源，包含常量、结构体、枚举定义，由双方共同维护，是客户端与服务端通信的契约。'
)

doc.add_heading('1. Config.h（游戏配置文件）', level=4)
doc.add_paragraph(
    '功能：定义游戏全局常量与坦克属性模板，统一游戏参数标准，避免硬编码。'
)
doc.add_paragraph(
    '核心结构体——TankConfig::Attrs：封装坦克基础属性，包含最大生命值(maxHP)、移动速度(speed)、'
    '旋转速度(rotateSpeed)、子弹伤害(damage)、子弹速度(bulletSpeed)、技能冷却(skillCooldown)、'
    '技能持续时间(skillDuration)。使用constexpr编译期常量定义三种坦克的预设属性，'
    '区分不同坦克的性能差异。'
)
doc.add_paragraph(
    '核心命名空间常量：TankConfig定义重型、轻型、侦察三种坦克的预设属性；'
    'MapConfig定义地图尺寸（宽1200px、高800px）、最大障碍物数量、以及默认地图布局函数'
    'GetDefaultObstacles()，配置了边界钢墙、中央对称掩体、水域阻隔、砖墙可破坏掩体、'
    '草丛隐身区域共20余个障碍物；CollisionConfig定义坦克碰撞半径(20px)和子弹碰撞半径(5px)；'
    'NetConfig定义网络收发缓冲区大小(4096B)。'
)
doc.add_paragraph(
    'Obstacle结构体：用(pos, size)描述障碍物的AABB包围盒，type区分砖墙/钢墙/草丛/水域四种类型，'
    'curHP/maxHP管理砖墙可破坏状态，destroyed标记是否已摧毁。'
)

doc.add_heading('2. Protocol.h（网络协议定义文件）', level=4)
doc.add_paragraph(
    '功能：定义客户端与服务端通信的统一协议，包含消息ID枚举、数据结构体、枚举类型，'
    '是整个联机通信的核心标准，任何修改必须双方确认。'
)
doc.add_paragraph('核心枚举：')
doc.add_paragraph(
    'MsgID：通信消息ID枚举，按功能分类——0x01xx认证类(C2S_LOGIN/C2S_REGISTER/'
    'S2C_LOGIN_ACK/S2C_REGISTER_ACK)、0x02xx房间类(C2S_JOIN_ROOM/C2S_LEAVE_ROOM/'
    'S2C_ROOM_INFO/S2C_MATCH_RESULT/C2S_SELECT_TANK/C2S_RECONNECT)、0x03xx游戏同步类'
    '(C2S_INPUT/S2C_SNAPSHOT/C2S_SHOOT/C2S_USE_SKILL/S2C_HIT/S2C_SKILL_EFFECT)、'
    '0x04xx结算类(S2C_GAME_OVER/C2S_GET_RANK/S2C_RANK_LIST)、0x05xx心跳类(C2S_HEARTBEAT/'
    'S2C_HEARTBEAT)。'
)
doc.add_paragraph(
    'TankType：坦克类型枚举（HEAVY重型/LIGHT轻型/SCOUT侦察）。'
)
doc.add_paragraph(
    'SkillType：技能类型枚举（SHIELD护盾/SPRINT冲刺/SCATTER散射）。'
)
doc.add_paragraph(
    'RoomState：房间状态枚举（WAITING等待/READY准备/PLAYING游戏中/PAUSED暂停）。'
)
doc.add_paragraph('核心结构体：')
doc.add_paragraph(
    'MsgHeader：消息头，包含MsgID消息类型与uint16_t消息体长度，用于TCP粘包处理时的消息边界识别。'
)
doc.add_paragraph(
    'TankState：坦克状态结构体，包含位置(pos)、角度(angle)、当前/最大血量(curHP/maxHP)、'
    '技能冷却/计时(skillCooldown/skillTimer)、类型(type)、玩家ID(playerID)、存活标志(alive)、'
    '护盾/冲刺激活标志(shieldActive/sprintActive)、射击冷却(shootTimer)，共14个字段。'
)
doc.add_paragraph(
    'BulletState：子弹状态结构体，包含位置(pos)、速度(vel)、发射者(owner)、伤害(damage)、'
    '类型(type)。'
)
doc.add_paragraph(
    'Snapshot：帧快照结构体，是服务端每Tick广播的权威游戏状态，包含帧序号(frameSeq)、'
    '2个TankState、64个BulletState槽位及当前数量、2个SkillEffectState、64个障碍物摧毁标志。'
    '单帧快照约2KB，在局域网环境下可高效传输。'
)
doc.add_paragraph(
    'InputState：玩家输入状态，封装WASD四个方向键、空格射击、F技能共6个bool标志。'
)

# ---------- (二) Core 核心模块 ----------
doc.add_heading('（二）Core 核心模块', level=3)
doc.add_paragraph(
    '本模块提供服务端的基础设施支持，包括高精度计时与日志记录。'
)

doc.add_heading('1. Clock（高精度时钟类）', level=4)
doc.add_paragraph(
    '功能：基于Windows QueryPerformanceCounter API实现微秒级高精度计时，'
    '为服务端的Tick循环、心跳检测、断线超时判定提供精确的时间基准。'
)
doc.add_paragraph(
    '核心成员函数：'
)
doc.add_paragraph(
    'Now()：静态方法，返回自系统启动以来的微秒级时间戳(int64_t)。'
    '内部使用static局部变量缓存QueryPerformanceFrequency结果，避免重复查询。'
)
doc.add_paragraph(
    'SecondsSince(int64_t startUs)：静态方法，计算从startUs时刻到当前经过的秒数(float)，'
    '用于心跳超时判定、断线重连窗口计时、对局时长统计等场景。'
)

doc.add_heading('2. Logger（日志系统类）', level=4)
doc.add_paragraph(
    '功能：单例模式的日志系统，提供线程安全的日志分级输出，同时输出到控制台和server.log文件，'
    '支持INFO/WARN/ERROR/GAME/CHEAT五种日志级别。'
)
doc.add_paragraph(
    '设计要点：采用C++11 Magic Static实现线程安全的单例；使用std::mutex保护文件写入避免多线程日志交错；'
    '日志格式为"[时间戳][级别] 消息内容"，时间戳精确到秒；析构时自动关闭文件句柄，确保日志完整落盘。'
)

# ---------- (三) Net 网络模块 ----------
doc.add_heading('（三）Net 网络模块', level=3)
doc.add_paragraph(
    '本模块是服务端的网络核心，负责TCP连接管理、客户端会话维护、消息接收与分发。'
)

doc.add_heading('1. TCPServer（TCP服务端类）', level=4)
doc.add_paragraph(
    '功能：服务端总控类，管理TCP监听、Accept连接循环、游戏Tick循环、房间列表、Session注册表。'
)
doc.add_paragraph('核心成员变量：')
doc.add_paragraph(
    'listenSocket：TCP监听套接字，绑定9527端口（可通过命令行参数自定义）。'
)
doc.add_paragraph(
    'clients：已连接客户端socket列表，由mutex保护。'
)
doc.add_paragraph(
    'rooms：房间列表，支持多房间同时游戏，由roomMutex保护。'
)
doc.add_paragraph(
    'sessionMap：playerID到Session*的映射表，用于按玩家ID发送消息，由sessionMapMutex保护。'
)
doc.add_paragraph(
    'tickThread：游戏Tick线程，以20Hz频率驱动所有活跃房间的逻辑更新。'
)
doc.add_paragraph('核心成员函数：')
doc.add_paragraph(
    'Init(port)：初始化Winsock2，创建监听socket，绑定端口，开始listen(backlog=100)。'
    '失败时输出WSA错误码便于排查。'
)
doc.add_paragraph(
    'Start()：将running标志置true，启动detached AcceptLoop线程和tickThread线程。'
)
doc.add_paragraph(
    'AcceptLoop()：使用select + 1秒超时的非阻塞模式循环accept新连接。'
    '每个新连接创建一个Session对象，分配playerID（以socket值标识），'
    '注册回调函数（onDisconnect/onJoinRoom/onLeaveRoom/onTryReconnect），'
    '然后启动该Session的接收线程。'
)
doc.add_paragraph(
    'GameTickLoop()：20Hz Tick主循环。每帧执行：(1)处理各房间的断线标记(disconnectPending)；'
    '(2)对PLAYING/PAUSED状态的房间执行Tick；(3)对PLAYING中的房间做心跳超时检测（3秒无心跳触发断线）；'
    '(4)清理空房间（WAITING状态下无玩家的房间）。每帧结束时计算耗时，sleep补齐至50ms间隔。'
)
doc.add_paragraph(
    'FindOrCreateRoom()：遍历房间列表，返回第一个未满房间；若无则新建Room。'
    '使用roomMutex保证线程安全。'
)
doc.add_paragraph(
    'HandleJoinRoom(playerID, username, tankType)：在roomMutex保护下原子完成Join+SelectTank，'
    '避免与GameTickLoop或其他玩家RecvThread产生竞态。设置Session的currentRoom指针，'
    '并回调room->sendToPlayer绑定SendToPlayer函数。'
)
doc.add_paragraph(
    'HandleReconnect(newPlayerID, username, session)：遍历所有房间，调用Room::TryReconnect。'
    '重连成功后更新Session注册、发送S2C_RECONNECT_ACK。'
)

doc.add_heading('2. Session（客户端会话类）', level=4)
doc.add_paragraph(
    '功能：管理单个客户端连接的生命周期，包括TCP消息收发、登录状态、房间关联、心跳时间戳。'
)
doc.add_paragraph('核心成员变量：')
doc.add_paragraph(
    'sock/recvBuf/recvLen：客户端socket、接收缓冲区(4096B)、当前缓冲数据长度，'
    '用于TCP粘包处理。'
)
doc.add_paragraph(
    'userName/isLoggedIn：用户名和登录状态标志，登录成功后设置。'
)
doc.add_paragraph(
    'playerID/currentRoom：玩家ID和当前所在房间指针。'
)
doc.add_paragraph(
    'lastHeartbeatUs：最后一次收到心跳的时间戳（微秒），服务端每3秒检查超时。'
)
doc.add_paragraph(
    'sendMutex：发送互斥锁，保证多线程（GameTickLoop广播 + RecvThread回复）发送安全。'
)
doc.add_paragraph('核心成员函数：')
doc.add_paragraph(
    'Send(MsgID, body, bodyLen)：循环发送确保TCP完整传输。先发送MsgHeader（4字节），'
    '再发送body。使用sendMutex保证线程安全。内部使用while循环处理TCP部分发送情况。'
)
doc.add_paragraph(
    'RecvThread(Session*)：独立线程函数，循环接收TCP数据。实现定长头+变长体协议解析：'
    '先读取4字节MsgHeader获取bodyLen，再读取bodyLen字节的消息体，调用ProcessMsg处理。'
    '若当前缓冲区不足以形成完整消息则等待更多数据到达。处理完一条消息后使用memmove移除已处理部分。'
    'recv返回<=0时退出循环，标记disconnectPending通知房间处理断线。'
)
doc.add_paragraph(
    'ProcessMsg(Session*, msgId, body, bodyLen)：消息分发函数，根据msgId路由到对应处理函数。'
    '支持完整的14种消息类型：注册(C2S_REGISTER→HandleRegister)、登录(C2S_LOGIN→HandleLogin)、'
    '加入/离开房间(C2S_JOIN_ROOM/C2S_LEAVE_ROOM)、选坦克(C2S_SELECT_TANK)、'
    '输入/射击/技能(C2S_INPUT/C2S_SHOOT/C2S_USE_SKILL→Room对应方法)、'
    '心跳(C2S_HEARTBEAT→更新lastHeartbeatUs并回复S2C_HEARTBEAT)、'
    '排行榜查询(C2S_GET_RANK→GetTopPlayers)、重连(C2S_RECONNECT→onTryReconnect)。'
    '登录状态下检查：未登录用户发送C2S_JOIN_ROOM时回复S2C_ERROR。'
)

# ---------- (四) Game 游戏模块 ----------
doc.add_heading('（四）Game 游戏模块', level=3)
doc.add_paragraph(
    '本模块是游戏逻辑的核心，实现房间管理、游戏世界Tick、权威判定、快照广播、断线处理。'
)

doc.add_heading('1. Room（房间管理类）', level=4)
doc.add_paragraph(
    '功能：管理一个游戏房间的完整生命周期，包括玩家加入/离开、坦克选择、游戏开始/结束、'
    '游戏Tick驱动、快照广播、断线重连。每个Room实例管理两个玩家slot（slot 0和slot 1）。'
)
doc.add_paragraph('核心成员变量：')
doc.add_paragraph(
    'state：房间状态机（WAITING→READY→PLAYING→PAUSED→WAITING循环）。'
    '两人加入后从WAITING变为READY；两人都选完坦克后StartGame变为PLAYING；'
    '断线触发PAUSED；游戏结束重置为WAITING。'
)
doc.add_paragraph(
    'playerIDs[2]/playerNames[2]/selectedTanks[2]/tankSelected[2]：分别存储两个玩家的ID、用户名、'
    '选择的坦克类型、是否已完成选择。'
)
doc.add_paragraph(
    'world：GameWorld实例，管理实际的游戏状态（坦克、子弹、障碍物）。'
)
doc.add_paragraph(
    'disconnectTimeUs[2]/lastHeartbeatUs[2]/disconnectPending[2]：断线管理数据。'
    'disconnectPending使用atomic<bool>，由RecvThread原子写入，GameTickLoop读取后处理。'
)
doc.add_paragraph(
    'sendToPlayer：函数回调，指向TCPServer::SendToPlayer，实现按playerID发送消息。'
)
doc.add_paragraph('核心成员函数：')
doc.add_paragraph(
    'Join(playerID, username)：向房间添加玩家。检查重复加入和房间是否已满，'
    '分配空闲slot，记录用户名和心跳时间。房间满两人时自动切换为READY状态。'
)
doc.add_paragraph(
    'SelectTank(playerID, type)：记录玩家选择的坦克类型，发送S2C_SELECT_TANK_ACK确认。'
    '禁止重复选择。两个玩家都选择完毕后自动调用StartGame()。'
)
doc.add_paragraph(
    'StartGame()：调用world.Init初始化两个坦克的起始位置和属性（玩家1在左(100,400)，'
    '玩家2在右(1100,400)），设置状态为PLAYING，记录gameStartUs，'
    '向双方发送S2C_MATCH_RESULT（包含双方playerID、坦克类型、起始位置）。'
)
doc.add_paragraph(
    'Tick(dt)：每50ms由GameTickLoop调用。在PLAYING状态：(1)调用world.Tick执行一个游戏帧；'
    '(2)将所有命中事件打包为S2C_HIT消息发送给双方；(3)调用BroadcastSnapshot广播快照；'
    '(4)检查world.gameOver，触发EndGame。在PAUSED状态：调用CheckDisconnectTimeout检查超时。'
)
doc.add_paragraph(
    'BroadcastSnapshot()：调用world.PackSnapshot()打包当前游戏状态，'
    '通过sendToPlayer回调向双方玩家发送S2C_SNAPSHOT消息。'
)
doc.add_paragraph(
    'HandleDisconnect(playerID)：断线处理核心逻辑。WAITING/READY状态下直接移除玩家。'
    'PLAYING状态下：记录disconnectTimeUs，状态切换为PAUSED，向双方广播暂停通知。'
    'PAUSED状态下（另一玩家也在暂停期间断开）：同样记录断开时间。'
)
doc.add_paragraph(
    'TryReconnect(username, newPlayerID)：遍历两个slot，通过用户名匹配找到断线玩家。'
    '更新playerID，清空disconnectTimeUs，恢复PLAYING状态。向重连客户端发送MatchResultData'
    '和当前Snapshot，确保重连后客户端能立即恢复游戏画面。'
)
doc.add_paragraph(
    'CheckDisconnectTimeout(dt)：检查断线时间是否超过RECONNECT_MS(10秒)，超时则调用ForfeitPlayer判负。'
)
doc.add_paragraph(
    'EndGame(winnerSlot, forfeit)：向双方发送S2C_GAME_OVER。保存对局记录（调用SaveGameRecord）、'
    '更新双方战绩（调用UpdateStats），然后重置房间状态（清空playerIDs、重置world）。'
)

doc.add_heading('2. GameWorld（游戏世界类）', level=4)
doc.add_paragraph(
    '功能：实现服务端权威游戏世界，管理坦克移动、子弹生命周期、碰撞检测与伤害计算、'
    '技能效果、障碍物交互、反作弊校验、快照打包。'
)
doc.add_paragraph('核心成员变量：')
doc.add_paragraph(
    'tanks[2]：双方坦克状态数组。'
)
doc.add_paragraph(
    'bullets[MAX_BULLETS]：子弹池，最多64颗同时飞行。使用紧凑数组管理，删除时用swap-and-pop避免碎片。'
)
doc.add_paragraph(
    'obstacles[64]：障碍物数组，服务端同样维护完整的地图状态，保证障碍物交互的一致性。'
)
doc.add_paragraph(
    'pendingInput[2]/shootRequested[2]/skillRequested[2]：输入缓冲，由inputMutex保护。'
    '客户端上报的输入先存入缓冲，Tick执行时统一消费。'
)
doc.add_paragraph(
    'lastValidPos[2]：上一帧各坦克的合法位置，用于反作弊校验——检测玩家是否通过修改客户端'
    '实现瞬移。'
)
doc.add_paragraph(
    'inputCountThisTick/shootSpamCount/skillSpamCount：反作弊统计计数器，记录单个Tick内'
    '的输入/射击/技能请求频率。'
)
doc.add_paragraph('核心成员函数：')
doc.add_paragraph(
    'Init(p1ID, t1, p2ID, t2)：初始化游戏世界。重置所有状态，根据TankType从TankConfig获取属性模板，'
    '设置双方坦克的起始位置、血量、朝向。调用MapConfig::GetDefaultObstacles初始化障碍物。'
)
doc.add_paragraph(
    'Tick(dt)：每帧主逻辑，执行顺序：(1)消费输入缓冲(shoot/skill请求)；'
    '(2)MoveTank移动双方坦克——根据WASD输入和朝向计算位移，应用冲刺速度加成；'
    '(3)更新技能冷却与射击冷却计时器；(4)处理射击请求SpawnBullet——在炮口前方生成子弹实例；'
    '(5)处理技能请求——护盾设置shieldActive，冲刺设置sprintActive，散射调用SpawnScatter发射3颗扇形弹；'
    '(6)DeactivateExpiredSkills——检查并关闭到期的护盾/冲刺效果；'
    '(7)MoveBullets移动所有子弹；(8)CheckBulletBounds出界销毁；'
    '(9)CheckBulletObstacleCollisions子弹-障碍物碰撞；'
    '(10)CheckBulletTankCollisions子弹-坦克碰撞并ApplyDamage；'
    '(11)检查坦克存活状态判定gameOver；(12)反作弊校验。'
)
doc.add_paragraph(
    'MoveTank(slot, dt)：根据pendingInput计算位移。A/D旋转，W/S前进/后退。'
    '冲刺激活时速度翻倍。移动后进行：(1)Judge::ClampToBounds边界钳制；'
    '(2)遍历所有障碍物做CircleRect碰撞检测，使用PushCircleOutOfRect推离穿模；'
    '(3)坦克间碰撞检测——两圆重叠时各推一半距离防止重叠；'
    '(4)再次边界钳制确保推开后不越界。'
)
doc.add_paragraph(
    'SpawnBullet(slot)：在坦克炮口前方（距离=TANK_RADIUS+BULLET_RADIUS+1）生成子弹，'
    '初速度方向=坦克朝向×子弹速度。'
)
doc.add_paragraph(
    'SpawnScatter(slot)：侦察坦克技能，在炮口前方以±15°开角生成3颗SCATTER类型子弹。'
)
doc.add_paragraph(
    'ApplyDamage(victimSlot, damage, attackerSlot)：伤害计算。护盾激活时伤害归零。'
    '扣血后记录HitEvent，血量降至0时标记死亡并增加击杀计数。'
)

doc.add_heading('3. Judge（碰撞判定工具命名空间）', level=4)
doc.add_paragraph(
    '功能：纯头文件实现的静态工具函数集合，提供碰撞检测与位置修正算法。'
)
doc.add_paragraph(
    'CircleCircle(aPos, aR, bPos, bR)：圆形-圆形碰撞检测，用于子弹-坦克命中判定。'
    '比较两圆心距离平方与半径和的平方。'
)
doc.add_paragraph(
    'CircleRect(circlePos, radius, rectPos, rectHalfSize)：圆-AABB碰撞检测，'
    '用于坦克/子弹与障碍物的碰撞。先计算圆心到矩形最近点的距离，再与半径比较。'
)
doc.add_paragraph(
    'PushCircleOutOfRect(circlePos, radius, rectPos, rectHalfSize)：碰撞响应函数，'
    '将圆形沿最近边法线方向推出矩形。用于坦克碰墙时的位置修正。'
)
doc.add_paragraph(
    'ClampToBounds(pos, radius)：边界钳制，确保实体不离开1200×800的地图区域。'
)
doc.add_paragraph(
    'InBounds(pos, radius)：出界检测，用于判断子弹是否飞离地图需销毁。'
)

# ---------- (五) Data 数据模块 ----------
doc.add_heading('（五）Data 数据模块', level=3)
doc.add_paragraph(
    '本模块实现用户数据持久化存储与查询，包括用户注册登录、战绩更新、对局记录、排行榜。'
)

doc.add_heading('1. UserManager（用户管理模块）', level=4)
doc.add_paragraph(
    '功能：基于文本文件的用户注册、登录验证、战绩更新，提供DJB2密码哈希存储。'
)
doc.add_paragraph(
    '数据格式：users.dat，每行"username|password_hash|wins|losses|kills"，'
    '以管道符分隔5个字段。'
)
doc.add_paragraph(
    '注册流程(Register)：逐行读取users.dat检查用户名唯一性，不存在则追加新行。'
    '密码通过DJB2哈希算法转为8位十六进制字符串存储，不保存明文。'
)
doc.add_paragraph(
    '登录流程(Login)：计算输入密码的哈希值，遍历文件匹配用户名+哈希值。'
    '匹配成功返回true。'
)
doc.add_paragraph(
    '战绩更新(UpdateStats)：使用临时文件users.tmp实现原子更新——'
    '读取原文件，逐行解析管道符分隔字段，找到目标用户后修改wins/losses/kills增量值，'
    '其余行原样写入。完成后remove+rename实现文件级别的原子替换。'
)
doc.add_paragraph(
    'GetAllUsers()：读取users.dat全部行，解析为vector<UserStats>供排行榜使用。'
)

doc.add_heading('2. RecordManager（对局记录模块）', level=4)
doc.add_paragraph(
    '功能：将对局结果持久化存储到records.dat文件。'
)
doc.add_paragraph(
    '数据格式：每行"时间戳|玩家1|玩家2|胜者|持续秒数|玩家1击杀数|玩家2击杀数"，'
    '7个字段以管道符分隔。使用追加模式写入，按时间顺序记录所有对局。'
)
doc.add_paragraph(
    'SaveGameRecord(rec)：打开records.dat追加写入一行，包括格式化时间戳和双方玩家数据。'
)

doc.add_heading('3. Leaderboard（排行榜模块）', level=4)
doc.add_paragraph(
    '功能：基于用户战绩数据生成排行榜。'
)
doc.add_paragraph(
    'GetTopPlayers(topN)：调用GetAllUsers获取全量用户数据，使用std::sort排序。'
    '排序规则：按胜率降序（胜率=(wins)/(wins+losses)），胜率相同时按击杀数降序。'
    '取前topN名，封装为RankListData结构体（最多10条RankEntry），内容包含用户名、胜场、败场、击杀数。'
)

# ---------- (六) main.cpp ----------
doc.add_heading('（六）main.cpp（服务端入口）', level=3)
doc.add_paragraph(
    '功能：服务端程序总入口，统筹所有模块的初始化与运行。核心流程：'
)
doc.add_paragraph(
    '1. 设置控制台输出编码为UTF-8（SetConsoleOutputCP(65001)），解决中文显示问题。'
)
doc.add_paragraph(
    '2. 解析命令行参数：支持自定义端口号"tank_server.exe [port]"，默认9527。'
)
doc.add_paragraph(
    '3. PrintLocalIPs()：通过GetAdaptersAddresses API枚举本机所有活跃网卡的IPv4地址，'
    '方便客户端通过局域网IP连接。跳过回环接口和未启用的网卡。'
)
doc.add_paragraph(
    '4. 创建TCPServer实例，调用Init→Start启动监听和Tick循环。'
)
doc.add_paragraph(
    '5. 主线程阻塞等待回车键输入，调用Shutdown清理所有资源：'
    '关闭running标志→join tickThread→关闭listenSocket→关闭所有client socket→'
    '删除所有Room→WSACleanup。'
)

# ============================================================
# 四、结果分析
# ============================================================
doc.add_heading('四、结果分析', level=2)

doc.add_heading('4.1 功能测试结果', level=3)

# Table for test results
table = doc.add_table(rows=12, cols=4)
table.style = 'Table Grid'
headers = ['测试项', '预期结果', '实际结果', '状态']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

test_data = [
    ['服务端启动', '监听9527端口，输出本机IP', '正常启动并输出IP列表', '通过'],
    ['客户端连接', 'Accept新连接，创建Session', '多客户端可同时连接', '通过'],
    ['用户注册', '账号写入users.dat，哈希存储密码', '正确存储并查重', '通过'],
    ['用户登录', '验证账号密码，更新登录态', '正确验证并设置标志', '通过'],
    ['房间匹配', '两人加入自动匹配，满员开始', '匹配逻辑正确', '通过'],
    ['快照广播', '20Hz广播完整游戏状态', '双客户端收到同步快照', '通过'],
    ['伤害判定', '子弹命中按属性扣血', '护盾免伤/普通伤害正确', '通过'],
    ['技能系统', '三种技能效果正确触发与到期', '护盾/冲刺/散射正常', '通过'],
    ['断线重连', '10秒内发送C2S_RECONNECT恢复', '状态正确恢复', '通过'],
    ['战绩记录', '对局结束写入records.dat', '数据格式正确', '通过'],
    ['排行榜', '按胜率排序返回前10名', '排序逻辑正确', '通过'],
]
for row_idx, row_data in enumerate(test_data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph('')

doc.add_heading('4.2 性能指标', level=3)
perf_data = [
    ['服务端Tick频率', '20Hz（50ms/帧）'],
    ['单帧快照大小', '约2KB（含2坦克+64子弹槽）'],
    ['单房间内存占用', '<1MB'],
    ['局域网延迟', '<5ms'],
    ['心跳超时阈值', '3秒'],
    ['重连等待窗口', '10秒'],
    ['最大并发子弹数', '64颗'],
    ['最大并发房间数', '理论上无限（受内存限制）'],
]
for label, value in perf_data:
    doc.add_paragraph(f'{label}：{value}')

doc.add_heading('4.3 反作弊机制', level=3)
doc.add_paragraph(
    '服务端实现了多层反作弊校验：(1)速度上限检查——每Tick计算坦克位移距离，'
    '超过理论最大速度（含冲刺加成+容差）则回退位置，防止客户端瞬移作弊；'
    '(2)输入频率监控——单个Tick内input超过5次、shoot超过3次、skill超过3次时记录作弊日志；'
    '(3)服务端权威判定——所有伤害计算、碰撞检测、技能触发均在服务端执行，客户端无法篡改游戏结果。'
)

doc.add_heading('4.4 不足与改进方向', level=3)
doc.add_paragraph(
    '1. 子弹速度单位不一致：初期客户端子弹用帧速移动而服务端用dt移动，已统一为基于dt的物理计算。'
)
doc.add_paragraph(
    '2. 协议结构体对齐：TankState曾缺少shootTimer字段导致sizeof不同、快照反序列化错位，已同步修复。'
)
doc.add_paragraph(
    '3. 用户数据明文存储：密码使用DJB2哈希但未加盐，可改用bcrypt等安全哈希。'
)
doc.add_paragraph(
    '4. 通信明文传输：网络数据未加密，可加入TLS或简单对称加密。'
)
doc.add_paragraph(
    '5. 无数据库：用户数据和战绩使用文本文件存储，大量用户场景下应改用SQLite等嵌入式数据库。'
)
doc.add_paragraph(
    '6. 线程模型可优化：当前每个连接一个recv线程，大量连接时可改用IOCP等异步IO模型。'
)

# ============================================================
# 五、课程设计总结
# ============================================================
doc.add_heading('五、课程设计总结', level=2)

doc.add_heading('5.1 技术收获', level=3)
doc.add_paragraph(
    '通过本次课程设计，我在以下方面获得了深入实践：'
)
doc.add_paragraph(
    '1. C++面向对象设计：通过服务端模块化架构的设计，深入理解了接口与实现分离、'
    '单例模式（Logger）、策略模式（TankConfig属性模板）的实际应用。不同坦克类型通过'
    'constexpr属性表而非虚函数实现差异化，兼顾了性能和扩展性。'
)
doc.add_paragraph(
    '2. 网络编程：掌握了WinSock2 TCP通信的完整流程，包括服务端listen/accept/bind、'
    'select非阻塞监听、定长消息头+变长消息体的粘包处理方案、多线程发送的互斥保护。'
    '理解了TCP可靠传输的特性——需要循环send确保完整发送——及优雅关闭连接的流程。'
)
doc.add_paragraph(
    '3. 多线程并发：服务端核心的并发场景包括：AcceptLoop线程、每个客户端一个RecvThread、'
    'GameTickLoop线程。使用mutex保护共享数据结构（clients列表、room列表、sessionMap、发送缓冲区），'
    '使用atomic<bool>传递断线信号避免RecvThread与Tick线程间的竞态。'
    '重点解决了Room::Join+SelectTank的原子性问题——在整个操作期间持有roomMutex防止与Tick线程交叉。'
)
doc.add_paragraph(
    '4. 实时同步架构：实现了完整的状态同步模式——服务端以20Hz频率执行权威Tick，'
    '广播包含帧序号的完整快照，客户端按帧序号丢弃旧包。设计了断线→暂停→重连→恢复的完整状态机，'
    '支持游戏中的无缝重连体验。'
)
doc.add_paragraph(
    '5. 游戏服务器设计：掌握了房间匹配、多房间管理、Tick调度、反作弊校验等游戏服务端核心机制。'
    '理解了服务端权威判定的重要性——所有游戏逻辑（移动、射击、技能、碰撞、伤害）由服务端计算，'
    '客户端仅做输入上报和画面渲染。'
)

doc.add_heading('5.2 团队协作经验', level=3)
doc.add_paragraph(
    '1. 协议先行：Common/Protocol.h是整个项目的契约，任何修改必须双方确认。'
    '项目中曾因TankState结构体字段不一致（客户端缺少shootTimer）导致sizeof(TankState)'
    '两端不同，快照反序列化时数据错位引发各种诡异bug。这个教训深刻体现了通信协议一致性的重要性。'
)
doc.add_paragraph(
    '2. 分支管理：采用dev/client和dev/server分支独立开发，定期合并到main。'
    '通过GitHub进行代码同步，避免了直接修改对方代码的混乱。'
)
doc.add_paragraph(
    '3. 联调流程：先各自开发并通过单元测试，再进行端到端联调。'
    '服务端先启动并输出本机IP，客户端配置IP后连接，按消息类型（先登录→再匹配→再游戏）逐步验证。'
)
doc.add_paragraph(
    '4. 接口约定：服务端通过sendToPlayer函数回调向客户端发送消息，而非直接访问Session，'
    '这种依赖倒置设计使得Room与Session解耦，便于单元测试和模块替换。'
)

doc.add_heading('5.3 个人体会', level=3)
doc.add_paragraph(
    '本次课程设计将课堂所学的C++编程、数据结构、操作系统、计算机网络等知识融会贯通，'
    '完成了一个从零到一可运行的网络对战游戏服务端。项目过程中遇到了诸多技术挑战：'
    'TCP粘包的协议解析、多线程间的竞态条件（disconnectPending的原子性设计）、'
    '协议结构体两端不一致导致的反序列化错位、MinGW下EasyX与UCRT的符号兼容问题等。'
    '通过查阅资料、分析日志、单步调试逐步解决，大幅提升了问题定位和系统调试能力。'
)
doc.add_paragraph(
    '双人联机坦克大战虽规模不大，但涵盖了网络游戏服务端的典型架构和核心机制——'
    '匹配、Tick、权威判定、快照同步、断线重连、数据持久化、排行榜。'
    '这次实践为今后从事游戏服务端开发或分布式系统开发打下了扎实的基础。'
)

# ============================================================
# 参考文献
# ============================================================
doc.add_heading('参考文献', level=1)
doc.add_paragraph('[1]《C++程序设计基础（上）》——周霭如 林伟健 编著')
doc.add_paragraph('[2] EasyX Help. https://easyx.cn/')
doc.add_paragraph('[3] Microsoft Docs - Windows Sockets 2. https://docs.microsoft.com/en-us/windows/win32/winsock/')
doc.add_paragraph('[4] Gaffer on Games - Fix Your Timestep! https://gafferongames.com/post/fix_your_timestep/')

# ============================================================
# SAVE
# ============================================================
output_path = r'C:\Users\Firefly\OneDrive\桌面\tank-war\服务端报告书_陈京锦.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
